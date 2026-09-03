import os, json, uuid, shutil
from datetime import datetime
from pathlib import Path
from flask import Flask, render_template, request, jsonify, send_file, abort
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR   = Path(__file__).parent
DATA_DIR   = BASE_DIR / "data"
OUTPUT_DIR = BASE_DIR / "output"
UPLOAD_DIR = BASE_DIR / "static" / "uploads"
for d in [DATA_DIR, OUTPUT_DIR, UPLOAD_DIR]: d.mkdir(parents=True, exist_ok=True)

ALLOWED = {"png","jpg","jpeg","gif","bmp","webp","heic","tiff"}
app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024
app.secret_key = "memoria-descriptiva-2024"

def allowed(f): return "." in f and f.rsplit(".",1)[1].lower() in ALLOWED
def load_form(fid):
    p = DATA_DIR / f"{fid}.json"
    if not p.exists(): return None
    return json.load(open(p, encoding="utf-8"))
def save_form(fid, data):
    data["updated_at"] = datetime.now().isoformat()
    json.dump(data, open(DATA_DIR/f"{fid}.json","w",encoding="utf-8"), ensure_ascii=False, indent=2)
def list_forms():
    forms=[]
    for p in sorted(DATA_DIR.glob("*.json"),key=os.path.getmtime,reverse=True):
        try:
            d=json.load(open(p,encoding="utf-8"))
            forms.append({"id":p.stem,"nombre":d.get("datos_generales",{}).get("nombre_tienda","Sin nombre"),
                "tipo_obra":d.get("datos_generales",{}).get("tipo_obra",""),
                "ciudad":d.get("datos_generales",{}).get("ciudad",""),
                "estado":d.get("estado","borrador"),"updated_at":d.get("updated_at","")})
        except: pass
    return forms

@app.route("/")
def index(): return render_template("index.html")
@app.route("/nuevo")
def nuevo():
    fid=str(uuid.uuid4())
    save_form(fid,{"id":fid,"estado":"borrador","created_at":datetime.now().isoformat()})
    return render_template("formulario.html", form_id=fid)
@app.route("/editar/<fid>")
def editar(fid):
    d=load_form(fid)
    if d is None: abort(404)
    return render_template("formulario.html", form_id=fid)
@app.route("/admin")
def admin(): return render_template("admin.html", forms=list_forms())

@app.route("/api/form/<fid>", methods=["GET"])
def api_get(fid):
    d=load_form(fid)
    if d is None: return jsonify({"error":"No encontrado"}),404
    return jsonify(d)
@app.route("/api/form/<fid>", methods=["POST"])
def api_save(fid):
    d=load_form(fid) or {"id":fid,"created_at":datetime.now().isoformat()}
    payload=request.get_json(force=True)
    d.update(payload)
    save_form(fid,d)
    return jsonify({"ok":True})
@app.route("/api/form/<fid>", methods=["DELETE"])
def api_delete(fid):
    p=DATA_DIR/f"{fid}.json"
    if p.exists(): p.unlink()
    up=UPLOAD_DIR/fid
    if up.exists(): shutil.rmtree(up)
    return jsonify({"ok":True})
@app.route("/api/upload/<fid>", methods=["POST"])
def api_upload(fid):
    if "file" not in request.files: return jsonify({"error":"No file"}),400
    file=request.files["file"]
    if not allowed(file.filename): return jsonify({"error":"Tipo no permitido"}),400
    dest=UPLOAD_DIR/fid; dest.mkdir(parents=True,exist_ok=True)
    ext=file.filename.rsplit(".",1)[1].lower()
    fname=f"{uuid.uuid4()}.{ext}"
    file.save(dest/fname)
    return jsonify({"url":f"/static/uploads/{fid}/{fname}","filename":fname})
@app.route("/api/upload/<fid>/<fname>", methods=["DELETE"])
def api_del_photo(fid,fname):
    fp=UPLOAD_DIR/fid/secure_filename(fname)
    if fp.exists(): fp.unlink()
    return jsonify({"ok":True})
@app.route("/api/generar/<fid>", methods=["POST"])
def api_generar(fid):
    d=load_form(fid)
    if d is None: return jsonify({"error":"No encontrado"}),404
    try:
        op=generar_word(d,fid)
        d["estado"]="completado"; d["word_path"]=str(op)
        save_form(fid,d)
        return jsonify({"ok":True,"download":f"/api/descargar/{fid}"})
    except Exception as e:
        import traceback
        return jsonify({"error":str(e),"trace":traceback.format_exc()}),500
@app.route("/api/descargar/<fid>")
def api_descargar(fid):
    d=load_form(fid)
    if d is None: abort(404)
    p=Path(d.get("word_path",""))
    if not p.exists(): abort(404)
    nombre=d.get("datos_generales",{}).get("nombre_tienda","Memoria") or "Memoria"
    return send_file(p,as_attachment=True,download_name=secure_filename(f"MEMORIA-{nombre}.docx"))

# ---- Word generator helpers ----
def add_heading(doc,text,level=1):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4)
    run=p.add_run(text.upper()); run.bold=True
    sizes={1:13,2:11,3:10}; colors={1:RGBColor(0,56,107),2:RGBColor(0,112,192),3:RGBColor(64,64,64)}
    run.font.size=Pt(sizes.get(level,10)); run.font.color.rgb=colors.get(level,RGBColor(0,0,0))
def add_field(doc,label,value):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    r1=p.add_run(f"{label}: "); r1.bold=True; r1.font.size=Pt(10)
    r2=p.add_run(str(value) if value else ""); r2.font.size=Pt(10)
def resolve_img_path(url):
    if not url: return None
    clean = url.lstrip("/").replace("/", os.sep)
    p1 = BASE_DIR / clean
    if p1.exists(): return p1
    p2 = BASE_DIR / "static" / clean
    if p2.exists(): return p2
    return None

def add_photo_grid(doc,photos,fid):
    if not photos: return
    for i in range(0,len(photos),2):
        pair=photos[i:i+2]
        table=doc.add_table(rows=2,cols=len(pair)); table.style="Table Grid"
        for ci,photo in enumerate(pair):
            ic=table.cell(0,ci); dc=table.cell(1,ci)
            url=photo.get("url","")
            img_path=resolve_img_path(url)
            if img_path and img_path.exists():
                try:
                    p=ic.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(str(img_path),width=Inches(2.8))
                except Exception as e:
                    print("Error inserting picture:", e)
                    ic.text="[Error al insertar imagen]"
            else:
                ic.text="[Imagen no disponible]"
            dc.text=photo.get("descripcion","")
        doc.add_paragraph()
def add_subsection(doc,titulo,contenido,photos,fid):
    add_heading(doc,titulo,2)
    if contenido and contenido.strip():
        p=doc.add_paragraph(contenido)
        for r in p.runs: r.font.size=Pt(10)
    add_photo_grid(doc,photos or [],fid)

def generar_word(data,fid):
    doc=Document()
    for s in doc.sections:
        s.top_margin=Cm(2); s.bottom_margin=Cm(2); s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)
    hdr=doc.sections[0].header
    hp=hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hp.clear(); hp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    hr=hp.add_run("MEMORIA DESCRIPTIVA"); hr.bold=True; hr.font.size=Pt(9); hr.font.color.rgb=RGBColor(128,128,128)
    dg=data.get("datos_generales",{})
    t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    tr=t.add_run(dg.get("nombre_tienda","MEMORIA DESCRIPTIVA").upper())
    tr.bold=True; tr.font.size=Pt(16); tr.font.color.rgb=RGBColor(0,56,107)
    doc.add_paragraph()
    add_heading(doc,"Datos Generales")
    for label,key in [("Tipo de Obra","tipo_obra"),("Calle","calle"),("Numero","numero"),("Colonia","colonia"),("Ciudad","ciudad"),("Estado","estado_rep"),("CP","cp"),("Sup. Construccion","sup_construccion"),("Sup. Remodelar","sup_remodelar"),("Sup. Total Predio","sup_total")]:
        add_field(doc,label,dg.get(key,""))
    doc.add_paragraph()
    add_heading(doc,"Croquis de Localizacion")
    cu=data.get("croquis",None)
    if cu and cu.get("url"):
        ip=resolve_img_path(cu["url"])
        if ip and ip.exists():
            try:
                p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(ip),width=Inches(4.5))
            except Exception as e:
                print("Error inserting croquis:", e)
                doc.add_paragraph("[Croquis no disponible]")
        else:
            doc.add_paragraph("[Croquis no disponible]")
    doc.add_paragraph()
    add_heading(doc,"Descripcion General")
    dgen=data.get("descripcion_general",{})
    if dgen.get("texto"):
        p=doc.add_paragraph(dgen["texto"])
        for r in p.runs: r.font.size=Pt(10)
    for item in dgen.get("lista_items",[]):
        if item and item.strip():
            p=doc.add_paragraph(style="List Bullet"); p.add_run(item).font.size=Pt(10)
    doc.add_paragraph()
    add_heading(doc,"Descripcion Remodelacion Exterior")
    ext=data.get("remodelacion_exterior",{})
    for key,lbl in [("cubierta","Cubierta"),("estacionamiento","Estacionamiento"),("anuncio_espectacular","Anuncio Espectacular"),("fachadas","Fachadas"),("anden","Anden"),("area_servicio","Area de Servicio")]:
        s=ext.get(key,{}); add_subsection(doc,lbl,s.get("texto",""),s.get("fotos",[]),fid)
    doc.add_paragraph()
    add_heading(doc,"Descripcion Remodelacion Interiores")
    intr=data.get("remodelacion_interiores",{})
    for key,lbl in [("portico_acceso","Portico de Acceso"),("oficinas_frontales","Oficinas Frontales"),("sanitarios_clientes","Sanitarios Clientes"),("piso_ventas","Piso de Ventas"),("area_perecederos","Area de Perecederos"),("comedor_asociados","Comedor de Asociados"),("oficinas_posteriores","Oficinas Posteriores"),("sanitarios_asociados","Sanitarios de Asociados"),("acceso_personal","Acceso de Personal"),("facturacion_sistemas","Facturacion y Sistemas"),("trastienda","Trastienda")]:
        s=intr.get(key,{}); add_subsection(doc,lbl,s.get("texto",""),s.get("fotos",[]),fid)
    doc.add_paragraph()
    add_heading(doc,"Giros de Negocio"); add_heading(doc,"Consultorio Medico",2)
    cm=data.get("giros_negocio",{}).get("consultorio_medico","")
    if cm:
        p=doc.add_paragraph(cm)
        for r in p.runs: r.font.size=Pt(10)
    doc.add_paragraph()
    add_heading(doc,"Descripcion de los Trabajos")
    tr2=data.get("descripcion_trabajos",{})
    if tr2.get("texto"):
        p=doc.add_paragraph(tr2["texto"])
        for r in p.runs: r.font.size=Pt(10)
    if tr2.get("area_trabajo"): add_field(doc,"Area de Trabajo",tr2["area_trabajo"])
    if tr2.get("afectacion_estructural"): add_field(doc,"Afectacion Estructural",tr2["afectacion_estructural"])
    doc.add_paragraph()
    add_heading(doc,"Estructura")
    et=data.get("estructura",{}).get("texto","")
    if et:
        p=doc.add_paragraph(et)
        for r in p.runs: r.font.size=Pt(10)
    doc.add_paragraph()
    add_heading(doc,"Instalaciones")
    inst=data.get("instalaciones",{})
    for key,lbl in [("refrigeracion","Sistema de Refrigeracion"),("aire","Sistema de Aire"),("electrica","Instalacion Electrica"),("hidro_sanitaria","Instalacion Hidro-Sanitaria"),("gas","Instalacion de Gas"),("filtrado","Sistema de Filtrado"),("contra_incendio","Instalacion Contra Incendio")]:
        t2=inst.get(key,"")
        if t2 and t2.strip():
            add_heading(doc,lbl,2)
            p=doc.add_paragraph(t2)
            for r in p.runs: r.font.size=Pt(10)
    doc.add_paragraph()
    add_heading(doc,"Medidas de Seguridad")
    seg=data.get("medidas_seguridad",{})
    for campo,lbl in [("texto",""),("medidas_clientes","Para Clientes y Asociados"),("consideraciones","Consideraciones Adicionales")]:
        v=seg.get(campo,"")
        if v and v.strip():
            if lbl: add_heading(doc,lbl,3)
            p=doc.add_paragraph(v)
            for r in p.runs: r.font.size=Pt(10)
    ftr=doc.sections[0].footer
    fp2=ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    fp2.clear(); fp2.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    fr2=fp2.add_run(f"Generado el {datetime.now().strftime(chr(37)+chr(100)+chr(47)+chr(37)+chr(109)+chr(47)+chr(37)+chr(89)+chr(32)+chr(37)+chr(72)+chr(58)+chr(37)+chr(77))}"); fr2.font.size=Pt(8); fr2.font.color.rgb=RGBColor(128,128,128)
    op=OUTPUT_DIR/f"MEMORIA-{fid[:8]}.docx"
    doc.save(str(op))
    return op

if __name__=="__main__":
    port = int(os.environ.get("PORT", 5000))
    print("="*50); print(f"  Memoria Descriptiva - Puerto {port}")
    print("="*50)
    app.run(debug=False, host="0.0.0.0", port=port)
